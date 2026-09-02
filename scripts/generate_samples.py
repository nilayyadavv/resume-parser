"""Create synthetic PDF and DOCX resumes under ./resumes."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from fpdf import FPDF

ROOT = Path(__file__).resolve().parents[1]
RESUMES = ROOT / "resumes"

SAMPLES: list[tuple[str, str, str]] = [
    (
        "maya-chen.pdf",
        "pdf",
        "\n".join(
            [
                "MAYA CHEN",
                "maya.chen@example.com",
                "https://github.com/mayachen",
                "https://www.linkedin.com/in/maya-chen",
                "Software engineer. Python, TypeScript, distributed systems.",
            ]
        ),
    ),
    (
        "jordan-patel.docx",
        "docx",
        "\n".join(
            [
                "Jordan Patel",
                "jordan.patel@example.com",
                "GitHub: github.com/jordanpatel",
                "Backend engineer focused on APIs and data pipelines.",
            ]
        ),
    ),
    (
        "sam-rivera.pdf",
        "pdf",
        "\n".join(
            [
                "Sam Rivera",
                "sam.rivera@example.net",
                "linkedin.com/in/samrivera",
                "Product designer. No GitHub listed.",
            ]
        ),
    ),
    (
        "alex-nguyen.docx",
        "docx",
        "\n".join(
            [
                "Alex Nguyen",
                "alex.nguyen@example.org",
                "Experience",
                "Worked on internal tools. Contact listed by email only.",
            ]
        ),
    ),
]


def write_pdf(path: Path, body: str) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for line in body.splitlines():
        pdf.cell(0, 8, line, new_x="LMARGIN", new_y="NEXT")
    pdf.output(str(path))


def write_docx(path: Path, body: str) -> None:
    document = Document()
    lines = body.splitlines()
    document.add_heading(lines[0], level=0)
    for line in lines[1:]:
        document.add_paragraph(line)
    document.save(str(path))


def main() -> None:
    RESUMES.mkdir(parents=True, exist_ok=True)
    for filename, kind, body in SAMPLES:
        path = RESUMES / filename
        if kind == "pdf":
            write_pdf(path, body)
        else:
            write_docx(path, body)
        print(f"wrote {path}")


if __name__ == "__main__":
    main()
