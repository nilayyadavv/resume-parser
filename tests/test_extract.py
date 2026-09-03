from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF

from resume_parser.extract import parse_path, parse_text


def test_extracts_name_email_github_linkedin():
    text = """
    MAYA CHEN
    maya.chen@example.com
    https://github.com/mayachen
    https://www.linkedin.com/in/maya-chen
    Software engineer
    """
    result = parse_text(text, filename="maya-chen.pdf")
    assert result.name == "Maya Chen"
    assert result.email == "maya.chen@example.com"
    assert result.github == "https://github.com/mayachen"
    assert result.linkedin == "https://www.linkedin.com/in/maya-chen"
    assert result.status == "ok"


def test_github_without_scheme_and_linkedin_pub():
    text = """
    Jordan Patel
    jordan.patel@example.com
    github.com/jordanpatel/dotfiles
    linkedin.com/pub/jordan-patel
    """
    result = parse_text(text, filename="jordan.docx")
    assert result.github == "https://github.com/jordanpatel"
    assert result.linkedin == "https://www.linkedin.com/pub/jordan-patel"


def test_skips_reserved_github_paths():
    text = """
    Sam Rivera
    sam.rivera@example.net
    https://github.com/features
    https://github.com/samrivera
    """
    result = parse_text(text, filename="sam.pdf")
    assert result.github == "https://github.com/samrivera"


def test_no_links_status_and_filename_fallback_name():
    text = "Contact me at alex.nguyen@example.org for roles in platform engineering."
    result = parse_text(text, filename="alex-nguyen-resume.docx")
    assert result.email == "alex.nguyen@example.org"
    assert result.github == ""
    assert result.linkedin == ""
    assert result.name == "Alex Nguyen"
    assert result.status == "no GitHub or LinkedIn found"


def test_empty_text():
    result = parse_text("   \n  ", filename="blank.pdf")
    assert result.status == "no extractable text"


def test_joins_split_first_and_last_name():
    text = "Nilay\n \nYadav\nnilayyadav2@vt.edu\nEDUCATION\nVirginia Tech"
    result = parse_text(text, filename="Thunder Resume.docx.pdf")
    assert result.name == "Nilay Yadav"
    assert result.email == "nilayyadav2@vt.edu"


def test_skips_placeholder_contact_and_generic_filename():
    text = "Full Name\nyourname@example.com\nwww.example.com\nEducation\nCollege"
    result = parse_text(text, filename="Basic professional resume.docx")
    assert result.email == ""
    assert result.name == ""
    assert result.github == ""
    assert result.linkedin == ""
    assert result.status == "no GitHub or LinkedIn found"


def test_reads_urls_split_across_lines_and_glued_labels():
    text = """
    Maya Chen
    maya.chen@example.com
    github.com/
    mayachen
    /linkedinlinkedin.com/in/maya-chen
    """
    result = parse_text(text, filename="maya.pdf")
    assert result.github == "https://github.com/mayachen"
    assert result.linkedin == "https://www.linkedin.com/in/maya-chen"


def test_skips_placeholder_github_and_linkedin_users():
    text = """
    Maya Chen
    maya.chen@example.com
    https://github.com/yourusername
    https://linkedin.com/in/yourusername
    https://github.com/mayachen
    """
    result = parse_text(text, filename="maya.pdf")
    assert result.github == "https://github.com/mayachen"
    assert result.linkedin == ""


def test_prefers_longer_linkedin_slug_over_truncated():
    text = """
    Maya Chen
    maya.chen@example.com
    linkedin.com/in/maya-che
    https://www.linkedin.com/in/maya-chen
    """
    result = parse_text(text, filename="maya.pdf")
    assert result.linkedin == "https://www.linkedin.com/in/maya-chen"


def test_country_linkedin_and_mailto():
    text = """
    Maya Chen
    mailto:maya.chen@example.com
    https://uk.linkedin.com/in/maya-chen
    """
    result = parse_text(text, filename="maya.pdf")
    assert result.email == "maya.chen@example.com"
    assert result.linkedin == "https://www.linkedin.com/in/maya-chen"


def test_generic_numbered_filename_is_not_used_as_name():
    result = parse_text("Skills\nPython, Java", filename="Resume (1).pdf")
    assert result.name == ""


def _add_docx_hyperlink(paragraph, label: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = label
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def test_pdf_hyperlink_annotations(tmp_path: Path):
    path = tmp_path / "linked.pdf"
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 8, "Maya Chen", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, "maya.chen@example.com", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 8, "GitHub", link="https://github.com/mayachen", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(
        0,
        8,
        "LinkedIn",
        link="https://www.linkedin.com/in/maya-chen",
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.output(str(path))
    result = parse_path(path)
    assert result.name == "Maya Chen"
    assert result.email == "maya.chen@example.com"
    assert result.github == "https://github.com/mayachen"
    assert result.linkedin == "https://www.linkedin.com/in/maya-chen"
    assert result.status == "ok"


def test_docx_hyperlinks_not_in_visible_text(tmp_path: Path):
    path = tmp_path / "linked.docx"
    document = Document()
    document.add_paragraph("Jordan Patel")
    document.add_paragraph("jordan.patel@example.com")
    paragraph = document.add_paragraph()
    _add_docx_hyperlink(paragraph, "GitHub", "https://github.com/jordanpatel")
    paragraph = document.add_paragraph()
    _add_docx_hyperlink(
        paragraph, "LinkedIn", "https://www.linkedin.com/in/jordan-patel"
    )
    document.save(str(path))
    result = parse_path(path)
    assert result.name == "Jordan Patel"
    assert result.github == "https://github.com/jordanpatel"
    assert result.linkedin == "https://www.linkedin.com/in/jordan-patel"
    assert result.status == "ok"


def test_docx_header_name(tmp_path: Path):
    path = tmp_path / "header.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "Sam Rivera"
    document.add_paragraph("sam.rivera@example.net")
    document.add_paragraph("github.com/samrivera")
    document.save(str(path))
    result = parse_path(path)
    assert result.name == "Sam Rivera"
    assert result.github == "https://github.com/samrivera"
