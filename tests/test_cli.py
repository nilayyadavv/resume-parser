from pathlib import Path

from docx import Document

from resume_parser.cli import main


def _write_docx(path: Path, lines: list[str]) -> None:
    document = Document()
    document.add_paragraph(lines[0])
    for line in lines[1:]:
        document.add_paragraph(line)
    document.save(str(path))


def test_cli_writes_csv(tmp_path: Path, capsys):
    folder = tmp_path / "resumes"
    folder.mkdir()
    _write_docx(
        folder / "maya-chen.docx",
        [
            "Maya Chen",
            "maya.chen@example.com",
            "https://github.com/mayachen",
            "https://www.linkedin.com/in/maya-chen",
        ],
    )
    csv_path = tmp_path / "out.csv"
    code = main([str(folder), "--csv", str(csv_path)])
    captured = capsys.readouterr()
    assert code == 0
    assert "wrote CSV" in captured.out
    content = csv_path.read_text(encoding="utf-8")
    assert "Maya Chen" in content
    assert "https://github.com/mayachen" in content


def test_cli_missing_folder(tmp_path: Path):
    code = main([str(tmp_path / "missing")])
    assert code == 2


def test_cli_empty_folder(tmp_path: Path):
    folder = tmp_path / "empty"
    folder.mkdir()
    code = main([str(folder)])
    assert code == 2


def test_cli_missing_credentials(tmp_path: Path):
    folder = tmp_path / "resumes"
    folder.mkdir()
    _write_docx(folder / "maya.docx", ["Maya Chen", "maya@example.com"])
    code = main(
        [
            str(folder),
            "--csv",
            str(tmp_path / "out.csv"),
            "--credentials",
            str(tmp_path / "nope.json"),
        ]
    )
    assert code == 2
